from pathlib import Path
import docx
import pytest

from awa.analysis.workflow_analyzer import analyze_workflow, analyze_canonical
from awa.model.workflow import Workflow, WorkflowMetadata
from awa.model.tool import Tool, Position, ToolConfiguration
from awa.model.connection import Connection
from awa.graph.builder import build_graph, execution_order
from awa.analysis.business_intelligence import generate_business_summary
from awa.generators.doc_builder import build_document_model
from awa.generators.docx_generator import generate_docx


class TestDocxBusinessReport:
    """Validate Executive Summary conforming to the professional Business Analyst assessment contract."""

    FORBIDDEN_USER_FACING_TERMS = [
        "Support Level",
        "FULL",
        "SUPPORTED",
        "PARTIAL",
        "PASS-THROUGH",
        "PASS_THROUGH",
        "DOCUMENTATION_ONLY",
        "EXTERNAL_EXECUTION",
        "UNSUPPORTED",
        "Analysis Diagnostics",
    ]

    def _extract_all_docx_text(self, docx_path: Path) -> str:
        doc = docx.Document(str(docx_path))
        text_parts = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                text_parts.append(" | ".join(c.text for c in row.cells))
        return "\n".join(text_parts)

    def _extract_executive_summary_text(self, docx_path: Path) -> str:
        doc = docx.Document(str(docx_path))
        in_exec = False
        exec_lines = []
        for p in doc.paragraphs:
            if p.text == "1. Executive Summary":
                in_exec = True
            elif p.text.startswith("2. "):
                in_exec = False
            if in_exec:
                exec_lines.append(p.text)
        return "\n".join(exec_lines)

    def test_simple_filter_docx_business_report(self, tmp_path: Path):
        """Test on a simple single-branch filter workflow."""
        out_dir = tmp_path / "simple_filter_report"
        analyze_workflow("fixtures/basic/simple_filter.yxmd", out_dir)
        docx_file = out_dir / "workflow.docx"
        assert docx_file.exists()

        full_text = self._extract_all_docx_text(docx_file)
        exec_text = self._extract_executive_summary_text(docx_file)

        # 1. No forbidden support classifications or diagnostic dumps
        for term in self.FORBIDDEN_USER_FACING_TERMS:
            assert term not in full_text, f"Forbidden term '{term}' found in simple_filter DOCX!"

        # 2. Executive Summary components present
        assert "1. Executive Summary" in exec_text
        assert "Methods of Analysis" in exec_text
        assert "Conclusions" in exec_text

        # 3. Concise length (100 to 350 words)
        words = len(exec_text.split())
        assert 50 <= words <= 350, f"Executive summary word count {words} out of expected range!"

        # 4. Report body sections present in Business Report
        assert "4. Source-to-Target Data Lineage" in full_text or "2. Business Process & Operational Deliverables" in full_text
        # 5. Omitted sections must be ABSENT from Business Report
        assert "Recommendations" not in full_text
        assert "Visual Workflow Graph" not in full_text
        assert "DAG Architecture" not in full_text
        assert "Step-by-Step Tool Specifications" not in full_text
        assert "Technical Configuration Appendix" not in full_text

    def test_join_workflow_docx_business_report(self, tmp_path: Path):
        """Test on a multi-input join workflow."""
        out_dir = tmp_path / "join_report"
        analyze_workflow("fixtures/joins/join_workflow.yxmd", out_dir)
        docx_file = out_dir / "workflow.docx"
        assert docx_file.exists()

        full_text = self._extract_all_docx_text(docx_file)
        exec_text = self._extract_executive_summary_text(docx_file)

        for term in self.FORBIDDEN_USER_FACING_TERMS:
            assert term not in full_text, f"Forbidden term '{term}' found in join_workflow DOCX!"

        assert "1. Executive Summary" in exec_text
        assert "Methods of Analysis" in exec_text
        assert "Conclusions" in exec_text

        words = len(exec_text.split())
        assert 50 <= words <= 350
        assert "Step-by-Step Tool Specifications" not in full_text
        assert "Technical Configuration Appendix" not in full_text

    def test_demo_claims_volume_extract_docx_business_report(self, tmp_path: Path):
        """Test on the full reconstructed Demo Claims regression fixture."""
        out_dir = tmp_path / "claims_report"
        analyze_workflow("Demo_Claims_Volume_Extract_reconstructed.yxmd", out_dir)
        docx_file = out_dir / "workflow.docx"
        assert docx_file.exists()

        full_text = self._extract_all_docx_text(docx_file)
        exec_text = self._extract_executive_summary_text(docx_file)

        # Verify Executive Summary headings
        assert "1. Executive Summary" in exec_text
        assert "Methods of Analysis" in exec_text
        assert "Findings" in exec_text
        assert "Conclusions" in exec_text
        assert "Recommendations" not in exec_text
        assert "Limitations" not in exec_text

        # Word count check (~100-350 words)
        words = len(exec_text.split())
        assert 100 <= words <= 350, f"Demo claims executive summary word count: {words}"

        # No raw tool IDs in Executive Summary
        assert "#1" not in exec_text
        assert "#39" not in exec_text

        # Verify body sections in Business Report
        assert "2. Business Process & Operational Deliverables" in full_text
        assert "2.1 Inputs & Upstream Dependencies" in full_text
        assert "2.2 Outputs & Business Reporting Deliverables" in full_text
        assert "2.3 Sequential Operational Stages" in full_text
        assert "3. Key Business Rules & Transformations" in full_text
        assert "4. Source-to-Target Data Lineage" in full_text
        assert "5. Visual Workflow Graph" not in full_text
        assert "DAG Architecture" not in full_text

        # Verify Technical sections are REMOVED from Business Report
        assert "Step-by-Step Tool Specifications" not in full_text
        assert "Technical Configuration Appendix" not in full_text

        # Verify specific business facts & real source filenames in body
        assert "Claims_Volume_Extract_Demo.xlsx" in full_text
        assert "Policy_Master_Demo.xlsx" in full_text
        assert "Claim_Payments_Demo.xlsx" in full_text
        assert "Claim_Diary_Notes_Demo.xlsx" in full_text
        assert "Claims_Volume_Extract_Demo.xlsx" in full_text
        assert "Output" in full_text or "Deliverables" in full_text

    def test_conditional_rendering_omits_empty_sections(self, tmp_path: Path):
        """Verify that when metadata/rules/outputs are missing, their headings are completely omitted."""
        xml_content = """<?xml version="1.0"?>
<AlteryxDocument yxmdVer="2024.1">
  <Nodes>
    <Node ToolID="1">
      <GuiSettings Plugin="AlteryxBasePluginsGui.DbFileInput.DbFileInput"><Position x="50" y="50" /></GuiSettings>
      <Properties><Configuration><File>raw_data.csv</File></Configuration></Properties>
    </Node>
    <Node ToolID="2">
      <GuiSettings Plugin="AlteryxBasePluginsGui.AlteryxSelect.AlteryxSelect"><Position x="150" y="50" /></GuiSettings>
      <Properties><Configuration /></Properties>
    </Node>
  </Nodes>
  <Connections>
    <Connection>
      <Origin ToolID="1" Connection="Output" />
      <Destination ToolID="2" Connection="Input" />
    </Connection>
  </Connections>
</AlteryxDocument>"""
        wf_file = tmp_path / "no_output.yxmd"
        wf_file.write_text(xml_content, encoding="utf-8")

        out_dir = tmp_path / "no_output_report"
        analyze_workflow(wf_file, out_dir)
        docx_file = out_dir / "workflow.docx"
        assert docx_file.exists()

        full_text = self._extract_all_docx_text(docx_file)

        # Because there are no business rules, Section 3 heading should NOT be rendered
        assert "3. Key Business Rules & Transformations" not in full_text

        # Because there are no outputs, Section 2.2 heading should NOT be rendered
        assert "2.2 Outputs & Business Reporting Deliverables" not in full_text

        # No empty headings or placeholder text
        assert "No information available" not in full_text
        assert "Business use not documented" not in full_text

    def test_download_docx_api_endpoint_contains_executive_summary(self):
        """Regression test: verify the FastAPI download endpoint returns populated Executive Summary in DOCX."""
        import io
        from fastapi.testclient import TestClient
        from backend.app.main import app

        client = TestClient(app)
        wf_path = Path("Demo_Claims_Volume_Extract_reconstructed.yxmd")
        with open(wf_path, "rb") as f:
            resp = client.post("/api/upload", files={"file": ("Demo_Claims.yxmd", f, "application/xml")})

        assert resp.status_code == 200
        analysis_id = resp.json()["analysis_id"]

        docx_resp = client.get(f"/api/download/{analysis_id}/docx")
        assert docx_resp.status_code == 200

        doc = docx.Document(io.BytesIO(docx_resp.content))
        headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
        assert "1. Executive Summary" in headings
        assert "5. Visual Workflow Graph (DAG Architecture)" not in headings

        # Verify text between "1. Executive Summary" and the next major heading is non-empty
        in_exec = False
        exec_lines = []
        for p in doc.paragraphs:
            if p.text == "1. Executive Summary":
                in_exec = True
            elif p.text.startswith("2. ") or p.text.startswith("3. ") or p.text.startswith("4. "):
                in_exec = False
            if in_exec:
                exec_lines.append(p.text)

        exec_text = "\n".join(exec_lines)
        assert len(exec_text.split()) >= 100
        assert "Methods of Analysis" in exec_text
        assert "Findings" in exec_text
        assert "Conclusions" in exec_text
        assert "Recommendations" not in exec_text
        assert "Limitations" not in exec_text

    def test_download_endpoints_for_both_reports(self):
        """Verify download endpoints return valid files for Business Report and Tool Specifications."""
        import io
        import openpyxl
        from fastapi.testclient import TestClient
        from backend.app.main import app

        client = TestClient(app)
        wf_path = Path("Demo_Claims_Volume_Extract_reconstructed.yxmd")
        with open(wf_path, "rb") as f:
            resp = client.post("/api/upload", files={"file": ("Demo_Claims.yxmd", f, "application/xml")})

        assert resp.status_code == 200
        analysis_id = resp.json()["analysis_id"]

        # 1. Business Report DOCX
        resp_biz = client.get(f"/api/download/{analysis_id}/docx")
        assert resp_biz.status_code == 200
        assert "Business_Report.docx" in resp_biz.headers.get("Content-Disposition", "")
        doc_biz = docx.Document(io.BytesIO(resp_biz.content))
        biz_headings = [p.text for p in doc_biz.paragraphs if p.style.name.startswith("Heading")]
        assert "1. Executive Summary" in biz_headings

        # 2. Tool Specifications XLSX
        resp_tool = client.get(f"/api/download/{analysis_id}/tool-specifications")
        assert resp_tool.status_code == 200
        assert "Tool_Specifications.xlsx" in resp_tool.headers.get("Content-Disposition", "")
        wb_tool = openpyxl.load_workbook(io.BytesIO(resp_tool.content))
        assert "Tool Specifications" in wb_tool.sheetnames

        # 3. ZIP bundle contains both
        resp_zip = client.get(f"/api/download/{analysis_id}/zip")
        assert resp_zip.status_code == 200
        import zipfile
        zf = zipfile.ZipFile(io.BytesIO(resp_zip.content))
        zip_names = zf.namelist()
        assert any("Business_Report.docx" in name for name in zip_names)
        assert any("Tool_Specifications.xlsx" in name for name in zip_names)
        assert not any("Technical_Specifications.docx" in name for name in zip_names)

    def test_canonical_filename_extraction_windows_and_unix_paths(self, tmp_path: Path):
        """Test 1 & 2: Given Windows or Unix path, canonical filename is extracted without directories."""
        xml_content = """<?xml version="1.0"?>
<AlteryxDocument yxmdVer="2024.1">
  <Nodes>
    <Node ToolID="1">
      <GuiSettings Plugin="AlteryxBasePluginsGui.DbFileInput.DbFileInput"><Position x="50" y="50" /></GuiSettings>
      <Properties><Configuration><File>C:\\Data\\Claim_Volume_Extract_Demo.xlsx|||Sheet1$</File></Configuration></Properties>
    </Node>
    <Node ToolID="2">
      <GuiSettings Plugin="AlteryxBasePluginsGui.DbFileInput.DbFileInput"><Position x="50" y="150" /></GuiSettings>
      <Properties><Configuration><File>/data/claims/Policy_Master_Unix.xlsx</File></Configuration></Properties>
    </Node>
  </Nodes>
  <Connections />
</AlteryxDocument>"""
        wf_file = tmp_path / "test_paths.yxmd"
        wf_file.write_text(xml_content, encoding="utf-8")

        canonical = analyze_canonical(wf_file)
        inputs = canonical.business_summary.source_inputs

        assert len(inputs) == 2
        inp1 = next(i for i in inputs if i.tool_id == 1)
        inp2 = next(i for i in inputs if i.tool_id == 2)

        assert inp1.source_filename == "Claim_Volume_Extract_Demo.xlsx"
        assert inp2.source_filename == "Policy_Master_Unix.xlsx"

    def test_annotation_not_authoritative_for_source_filename(self, tmp_path: Path):
        """Test 3: Annotation text must not override or determine the real filename from XML config."""
        xml_content = """<?xml version="1.0"?>
<AlteryxDocument yxmdVer="2024.1">
  <Nodes>
    <Node ToolID="1">
      <GuiSettings Plugin="AlteryxBasePluginsGui.DbFileInput.DbFileInput"><Position x="50" y="50" /></GuiSettings>
      <Properties>
        <Configuration><File>C:\\Data\\Claim_Volume_Extract_Demo.xlsx|||Sheet1$</File></Configuration>
        <Annotation DisplayMode="0"><DefaultAnnotationText>Claims Volume</DefaultAnnotationText></Annotation>
      </Properties>
    </Node>
  </Nodes>
  <Connections />
</AlteryxDocument>"""
        wf_file = tmp_path / "test_annotation.yxmd"
        wf_file.write_text(xml_content, encoding="utf-8")

        canonical = analyze_canonical(wf_file)
        inp = canonical.business_summary.source_inputs[0]

        assert inp.source_filename == "Claim_Volume_Extract_Demo.xlsx"

    def test_overview_dto_contains_canonical_source_filenames(self):
        """Test 4: Overview DTO for Demo Claims workflow contains canonical source filenames."""
        from fastapi.testclient import TestClient
        from backend.app.main import app

        client = TestClient(app)
        wf_path = Path("Demo_Claims_Volume_Extract_reconstructed.yxmd")
        with open(wf_path, "rb") as f:
            resp = client.post("/api/upload", files={"file": ("Demo_Claims.yxmd", f, "application/xml")})

        assert resp.status_code == 200
        data = resp.json()
        source_inputs = data["business_summary"]["source_inputs"]

        filenames = [si["source_filename"] for si in source_inputs]
        assert "Claims_Volume_Extract_Demo.xlsx" in filenames
        assert "Policy_Master_Demo.xlsx" in filenames
        assert "Claim_Payments_Demo.xlsx" in filenames
        assert "Claim_Diary_Notes_Demo.xlsx" in filenames

    def test_fully_llm_authored_business_report_integration(self, tmp_path: Path):
        """Test full LLM-authored report generation using MockLLMClient."""
        import json
        from awa.llm.client import MockLLMClient, set_default_llm_client
        from awa.llm.generator import LLMNarrativeGenerator, set_default_generator
        from awa.generators.doc_builder import build_document_model
        from awa.generators.docx_generator import generate_docx

        mock_report = {
            "workflow_title": "Automated Customer Sales Reporting Pipeline",
            "workflow_description": "Production ETL pipeline consolidating customer transactions",
            "executive_summary": "This automated data workflow ingests daily customer transactions and generates aggregated portfolio revenue summaries for management reporting.",
            "methods_of_analysis": "The process joins customer master attributes with transaction event records, applies business calculation formulas to compute tax and net amounts, filters active accounts, and aggregates metrics by product line.",
            "findings": [
                "Combines two primary enterprise datasets into a unified analytical base.",
                "Enforces data standardization through automated calculation and zero-fill rules.",
                "Branches transformed data into specialized reporting extracts.",
                "Relies on external file storage for source ingestion."
            ],
            "conclusions": "The workflow serves as a centralized transaction preparation pipeline delivering reconciled revenue metrics to business stakeholders.",
            "inputs": [
                {
                    "source_dataset": "customers.xlsx",
                    "business_role": "Customer master dimension table",
                    "source_format": "Excel Workbook",
                    "dependency_significance": "Mandatory primary reference dataset"
                },
                {
                    "source_dataset": "orders.csv",
                    "business_role": "Transactional order history feed",
                    "source_format": "CSV Data File",
                    "dependency_significance": "Daily transactional volume source"
                }
            ],
            "outputs": [
                {
                    "output_deliverable": "customer_totals.xlsx",
                    "what_it_represents": "Reconciled customer sales revenue summary",
                    "business_use": "Monthly commercial revenue audit",
                    "destination_format": "Excel Workbook"
                }
            ],
            "sequential_stages": [
                {
                    "stage_number": 1,
                    "stage_name": "Source Ingestion",
                    "description": "Loads customer master and order transaction files",
                    "operational_explanation": "Validates schema and parses columnar datatypes"
                },
                {
                    "stage_number": 2,
                    "stage_name": "Relational Join & Aggregation",
                    "description": "Joins orders to customer records and calculates totals",
                    "operational_explanation": "Executes inner join on customer_id and computes sum of amounts"
                },
                {
                    "stage_number": 3,
                    "stage_name": "Deliverable Export",
                    "description": "Exports finalized customer totals to Excel",
                    "operational_explanation": "Formats numeric columns and writes workbook"
                }
            ],
            "business_rules": [
                {
                    "business_rule": "Customer ID Reconciliation",
                    "category": "Integration",
                    "evidence_configuration": "customers.customer_id = orders.customer_id"
                },
                {
                    "business_rule": "Revenue Total Rollup",
                    "category": "Aggregation",
                    "evidence_configuration": "SUM(orders.amount) GROUP BY customer_id, name"
                }
            ],
            "lineage": [
                {
                    "source_datasets": "customers.xlsx + orders.csv",
                    "major_business_transformation": "Join customer profiles with orders, aggregate revenue, and sort descending",
                    "target_deliverable": "customer_totals.xlsx"
                }
            ]
        }

        mock_client = MockLLMClient(default_response=json.dumps(mock_report))
        mock_gen = LLMNarrativeGenerator(client=mock_client)
        set_default_llm_client(mock_client)
        set_default_generator(mock_gen)

        try:
            canonical = analyze_canonical("fixtures/joins/join_workflow.yxmd")
            doc_model = build_document_model(
                canonical.workflow,
                canonical.execution_order,
                canonical.translations,
                canonical.dag_layout,
                canonical.lineage_paths,
                business_summary=canonical.business_summary,
                analysis_id=canonical.analysis_id,
                graph=canonical.graph,
            )

            out_docx = tmp_path / "llm_business_report.docx"
            generate_docx(doc_model, out_docx)
            assert out_docx.exists()

            full_text = self._extract_all_docx_text(out_docx)
            exec_text = self._extract_executive_summary_text(out_docx)

            # 1. Section 1 Executive Summary populated from LLM
            assert "1. Executive Summary" in exec_text
            assert "Methods of Analysis" in exec_text
            assert "Findings" in exec_text
            assert "Conclusions" in exec_text
            assert "This automated data workflow ingests daily customer transactions" in exec_text
            assert "Combines two primary enterprise datasets into a unified analytical base" in exec_text
            assert "The workflow serves as a centralized transaction preparation pipeline" in exec_text

            # 2. Section 2 Inputs & Outputs populated from LLM
            assert "2. Business Process & Operational Deliverables" in full_text
            assert "2.1 Inputs & Upstream Dependencies" in full_text
            assert "customers.xlsx" in full_text
            assert "Customer master dimension table" in full_text
            assert "orders.csv" in full_text
            assert "Transactional order history feed" in full_text

            assert "2.2 Outputs & Business Reporting Deliverables" in full_text
            assert "customer_totals.xlsx" in full_text
            assert "Reconciled customer sales revenue summary" in full_text
            assert "Monthly commercial revenue audit" in full_text

            assert "2.3 Sequential Operational Stages" in full_text
            assert "Source Ingestion" in full_text
            assert "Relational Join & Aggregation" in full_text

            # 3. Section 3 Business Rules populated from LLM
            assert "3. Key Business Rules & Transformations" in full_text
            assert "Customer ID Reconciliation" in full_text
            assert "Revenue Total Rollup" in full_text

            # 4. Section 4 Lineage populated from LLM
            assert "4. Source-to-Target Data Lineage" in full_text
            assert "Join customer profiles with orders" in full_text

            # 5. Omitted sections are NOT present
            assert "Recommendations" not in full_text
            assert "Visual Workflow Graph" not in full_text
            assert "DAG Architecture" not in full_text
        finally:
            set_default_llm_client(None)
            set_default_generator(None)

    def test_no_claims_hardcoding_in_generic_uncontainerized_workflow(self, tmp_path: Path):
        """Verify that a generic workflow with no containers does NOT produce claims-demo stage names or descriptions."""
        wf = Workflow(
            metadata=WorkflowMetadata(name="Sales Inventory Sync", version="2023.2"),
            tools={
                1: Tool(
                    tool_id=1,
                    plugin="AlteryxBasePluginsGui.DbFileInput.DbFileInput",
                    tool_type="DbFileInput",
                    name="Input Inventory",
                    position=Position(0, 0),
                    configuration=ToolConfiguration(raw_xml="<Configuration/>", parsed={"file_path": "inventory.csv"}),
                ),
                2: Tool(
                    tool_id=2,
                    plugin="AlteryxBasePluginsGui.Filter.Filter",
                    tool_type="Filter",
                    name="Filter Active",
                    position=Position(100, 0),
                    configuration=ToolConfiguration(raw_xml="<Configuration/>", parsed={"Expression": "[Active] == 1"}),
                ),
                3: Tool(
                    tool_id=3,
                    plugin="AlteryxBasePluginsGui.DbFileOutput.DbFileOutput",
                    tool_type="DbFileOutput",
                    name="Output Active Inventory",
                    position=Position(200, 0),
                    configuration=ToolConfiguration(raw_xml="<Configuration/>", parsed={"file_path": "active_inventory.csv"}),
                ),
            },
            connections=[
                Connection(origin_tool_id=1, origin_anchor="Output", destination_tool_id=2, destination_anchor="Input"),
                Connection(origin_tool_id=2, origin_anchor="True", destination_tool_id=3, destination_anchor="Input"),
            ],
        )

        graph = build_graph(wf)
        exec_order = execution_order(graph)
        bs = generate_business_summary(wf, graph, exec_order)

        doc_model = build_document_model(
            workflow=wf,
            execution_order=exec_order,
            translations={},
            dag_layout=None,
            lineage_paths=[],
            business_summary=bs,
            analysis_id="generic-test-1",
        )

        out_docx = tmp_path / "generic_report.docx"
        generate_docx(doc_model, out_docx)
        assert out_docx.exists()

        full_text = self._extract_all_docx_text(out_docx)

        # Prohibit claims-specific legacy stage names and descriptions
        forbidden_strings = [
            "Historical volume and team aggregations",
            "Cross-source data enrichment and calculations",
            "Publish analytical deliverables",
            "Business use: Not documented",
            "Recommendations",
            "Limitations",
            "Visual Workflow Graph",
            "DAG Architecture",
            "Section 5",
        ]
        for s in forbidden_strings:
            assert s not in full_text, f"Found forbidden legacy string in generic report: {s}"

    def test_multi_domain_analytical_differentiation_and_no_leakage(self, tmp_path: Path):
        """Verify that distinct business domains (Banking vs Commercial Sales) produce distinct analytical narratives with zero cross-domain leakage and zero excluded sections."""
        from awa.llm.client import FakeLLMClient, set_default_llm_client
        from awa.llm.generator import LLMNarrativeGenerator, set_default_generator
        from awa.llm.cache import LLMNarrativeCache
        import json

        # Domain 1: Banking Transaction Settlement Workflow
        banking_wf = Workflow(
            metadata=WorkflowMetadata(name="Core Banking Settlement Analysis", version="2024.1"),
            tools={
                101: Tool(
                    tool_id=101,
                    plugin="AlteryxBasePluginsGui.DbFileInput.DbFileInput",
                    tool_type="DbFileInput",
                    name="Input Wire Transfers",
                    position=Position(0, 0),
                    configuration=ToolConfiguration(raw_xml="<Configuration/>", parsed={"file_path": "wire_transfers.csv"}),
                ),
                102: Tool(
                    tool_id=102,
                    plugin="AlteryxBasePluginsGui.Summarize.Summarize",
                    tool_type="Summarize",
                    name="Summarize Volume",
                    position=Position(100, 0),
                    configuration=ToolConfiguration(
                        raw_xml="<Configuration/>",
                        parsed={"summarize_fields": [{"field": "Transfer_Amount", "action": "Sum"}, {"field": "Account_ID", "action": "CountDistinct"}]},
                    ),
                ),
                103: Tool(
                    tool_id=103,
                    plugin="AlteryxBasePluginsGui.DbFileOutput.DbFileOutput",
                    tool_type="DbFileOutput",
                    name="Output Daily Settlement",
                    position=Position(200, 0),
                    configuration=ToolConfiguration(raw_xml="<Configuration/>", parsed={"file_path": "daily_settlement.xlsx"}),
                ),
            },
            connections=[
                Connection(origin_tool_id=101, origin_anchor="Output", destination_tool_id=102, destination_anchor="Input"),
                Connection(origin_tool_id=102, origin_anchor="Output", destination_tool_id=103, destination_anchor="Input"),
            ],
        )

        # Domain 2: Commercial Sales & Commission Workflow
        sales_wf = Workflow(
            metadata=WorkflowMetadata(name="Commercial Sales Commission Reconciliation", version="2024.1"),
            tools={
                201: Tool(
                    tool_id=201,
                    plugin="AlteryxBasePluginsGui.DbFileInput.DbFileInput",
                    tool_type="DbFileInput",
                    name="Input Sales Invoices",
                    position=Position(0, 0),
                    configuration=ToolConfiguration(raw_xml="<Configuration/>", parsed={"file_path": "sales_invoices.xlsx"}),
                ),
                202: Tool(
                    tool_id=202,
                    plugin="AlteryxBasePluginsGui.Formula.Formula",
                    tool_type="Formula",
                    name="Calculate Commission",
                    position=Position(100, 0),
                    configuration=ToolConfiguration(
                        raw_xml="<Configuration/>",
                        parsed={"formula_fields": [{"field": "Commission_Amount", "expression": "[Invoice_Total] * 0.05"}]},
                    ),
                ),
                203: Tool(
                    tool_id=203,
                    plugin="AlteryxBasePluginsGui.DbFileOutput.DbFileOutput",
                    tool_type="DbFileOutput",
                    name="Output Sales Ledger",
                    position=Position(200, 0),
                    configuration=ToolConfiguration(raw_xml="<Configuration/>", parsed={"file_path": "sales_commission_ledger.csv"}),
                ),
            },
            connections=[
                Connection(origin_tool_id=201, origin_anchor="Output", destination_tool_id=202, destination_anchor="Input"),
                Connection(origin_tool_id=202, origin_anchor="Output", destination_tool_id=203, destination_anchor="Input"),
            ],
        )

        def mock_domain_llm(system: str, user: str) -> str | None:
            if "wire_transfers.csv" in user or "Banking Settlement" in user:
                return json.dumps({
                    "workflow_title": "Core Banking Settlement Analysis",
                    "workflow_description": "Automates interbank wire transfer validation, ledger reconciliation, and daily liquidity position reporting.",
                    "executive_summary": "This analytical report evaluates the daily wire transfer settlement process across institutional banking counterparties. The workflow ingests raw wire transaction logs, computes total settlement exposure, calculates distinct active accounts, and outputs finalized ledger balances for central bank reconciliation.",
                    "methods_of_analysis": "The analysis utilizes multi-dimensional summation (SUM on Transfer_Amount) and distinct population enumeration (COUNT DISTINCT on Account_ID) to establish aggregate liquidity metrics across account hierarchies.",
                    "findings": [
                        "Interbank transfer volume is aggregated using SUM(Transfer_Amount) to calculate gross daily settlement exposure across active account partitions.",
                        "Counterparty participation is quantified using COUNT DISTINCT(Account_ID), establishing the total unique institutional account population per settlement window.",
                        "Finalized liquidity positions are distributed to daily_settlement.xlsx to support regulatory capital reporting and daily cash clearing audits."
                    ],
                    "conclusions": "The workflow establishes an authoritative daily settlement baseline, converting granular wire events into consolidated counterparty liquidity measures required for financial auditability.",
                    "inputs": [
                        {
                            "source_dataset": "wire_transfers.csv",
                            "business_role": "Inbound interbank wire transfer log feed",
                            "source_format": "CSV Data File",
                            "dependency_significance": "Essential transactional input required for central ledger reconciliation"
                        }
                    ],
                    "outputs": [
                        {
                            "output_deliverable": "daily_settlement.xlsx",
                            "what_it_represents": "Finalized institutional settlement position matrix",
                            "business_use": "Daily regulatory liquidity reporting and central bank clearing verification",
                            "destination_format": "Excel Workbook"
                        }
                    ],
                    "sequential_stages": [
                        {
                            "stage_number": 1,
                            "stage_name": "Transaction Ingestion & Aggregation",
                            "description": "Ingests raw wire transfer events and computes gross settlement exposure",
                            "operational_explanation": "Processes wire transfer records via aggregation"
                        }
                    ],
                    "business_rules": [
                        {
                            "business_rule": "Aggregate Settlement Exposure",
                            "category": "Aggregation",
                            "evidence_configuration": "SUM(Transfer_Amount), COUNT DISTINCT(Account_ID)"
                        }
                    ],
                    "lineage": [
                        {
                            "source_datasets": "wire_transfers.csv",
                            "major_business_transformation": "Aggregate transfer totals and count distinct accounts",
                            "target_deliverable": "daily_settlement.xlsx"
                        }
                    ]
                })
            elif "sales_invoices.xlsx" in user or "Sales Commission" in user:
                return json.dumps({
                    "workflow_title": "Commercial Sales Commission Reconciliation",
                    "workflow_description": "Computes sales team commission accruals from completed customer invoices.",
                    "executive_summary": "This business report evaluates commercial sales revenue and associate commission calculation logic. The workflow ingests finalized customer sales invoices, calculates contractual 5% commission amounts using derived formulas, and generates the monthly commission reconciliation ledger.",
                    "methods_of_analysis": "The workflow applies mathematical formula expressions ([Invoice_Total] * 0.05) to derive associate payout metrics from transaction-level invoice data.",
                    "findings": [
                        "Contractual commission accruals are calculated using deterministic formula derivation ([Invoice_Total] * 0.05), ensuring standard commercial rate application across all completed invoices.",
                        "Inbound invoice data is validated against sales ledger constraints prior to commission ledger posting.",
                        "The resulting commission ledger is published to sales_commission_ledger.csv for monthly payroll audit and revenue operations review."
                    ],
                    "conclusions": "The workflow establishes a standardized commercial commission derivation pipeline, transforming raw invoice data into audited compensation records.",
                    "inputs": [
                        {
                            "source_dataset": "sales_invoices.xlsx",
                            "business_role": "Completed customer sales invoice register",
                            "source_format": "Excel Workbook",
                            "dependency_significance": "Primary revenue source feed required for commission calculation"
                        }
                    ],
                    "outputs": [
                        {
                            "output_deliverable": "sales_commission_ledger.csv",
                            "what_it_represents": "Itemized sales representative commission register",
                            "business_use": "Monthly commercial sales payroll processing and revenue audit",
                            "destination_format": "CSV Data File"
                        }
                    ],
                    "sequential_stages": [
                        {
                            "stage_number": 1,
                            "stage_name": "Invoice Ingestion & Commission Derivation",
                            "description": "Ingests sales invoices and computes commission accruals",
                            "operational_explanation": "Applies formula derivation to invoices"
                        }
                    ],
                    "business_rules": [
                        {
                            "business_rule": "Calculate 5% Commission",
                            "category": "Calculation",
                            "evidence_configuration": "[Invoice_Total] * 0.05"
                        }
                    ],
                    "lineage": [
                        {
                            "source_datasets": "sales_invoices.xlsx",
                            "major_business_transformation": "Calculate commission accruals from invoice totals",
                            "target_deliverable": "sales_commission_ledger.csv"
                        }
                    ]
                })
            return None

        # Domain 3: Insurance Claims & Loss Analysis Workflow
        claims_wf = Workflow(
            metadata=WorkflowMetadata(name="Insurance Claims Loss Frequency Analysis", version="2024.1"),
            tools={
                301: Tool(
                    tool_id=301,
                    plugin="AlteryxBasePluginsGui.DbFileInput.DbFileInput",
                    tool_type="DbFileInput",
                    name="Input Policy Claims",
                    position=Position(0, 0),
                    configuration=ToolConfiguration(raw_xml="<Configuration/>", parsed={"file_path": "policy_claims_register.xlsx"}),
                ),
                302: Tool(
                    tool_id=302,
                    plugin="AlteryxBasePluginsGui.Summarize.Summarize",
                    tool_type="Summarize",
                    name="Aggregate Paid Losses",
                    position=Position(100, 0),
                    configuration=ToolConfiguration(
                        raw_xml="<Configuration/>",
                        parsed={"summarize_fields": [{"field": "Incurred_Loss", "action": "Sum"}, {"field": "Policy_Number", "action": "CountDistinct"}]},
                    ),
                ),
                303: Tool(
                    tool_id=303,
                    plugin="AlteryxBasePluginsGui.DbFileOutput.DbFileOutput",
                    tool_type="DbFileOutput",
                    name="Output Loss Ratio Matrix",
                    position=Position(200, 0),
                    configuration=ToolConfiguration(raw_xml="<Configuration/>", parsed={"file_path": "quarterly_actuarial_losses.xlsx"}),
                ),
            },
            connections=[
                Connection(origin_tool_id=301, origin_anchor="Output", destination_tool_id=302, destination_anchor="Input"),
                Connection(origin_tool_id=302, origin_anchor="Output", destination_tool_id=303, destination_anchor="Input"),
            ],
        )

        def mock_domain_llm(system: str, user: str) -> str | None:
            if "wire_transfers.csv" in user or "Banking Settlement" in user:
                return json.dumps({
                    "workflow_title": "Core Banking Settlement Analysis",
                    "workflow_description": "Automates interbank wire transfer validation, ledger reconciliation, and daily liquidity position reporting.",
                    "executive_summary": "This analytical report evaluates the daily wire transfer settlement process across institutional banking counterparties. The workflow ingests raw wire transaction logs, computes total settlement exposure, calculates distinct active accounts, and outputs finalized ledger balances for central bank reconciliation.",
                    "methods_of_analysis": "The analysis utilizes multi-dimensional summation (SUM on Transfer_Amount) and distinct population enumeration (COUNT DISTINCT on Account_ID) to establish aggregate liquidity metrics across account hierarchies.",
                    "findings": [
                        "Interbank transfer volume is aggregated using SUM(Transfer_Amount) to calculate gross daily settlement exposure across active account partitions.",
                        "Counterparty participation is quantified using COUNT DISTINCT(Account_ID), establishing the total unique institutional account population per settlement window.",
                        "Finalized liquidity positions are distributed to daily_settlement.xlsx to support regulatory capital reporting and daily cash clearing audits."
                    ],
                    "conclusions": "The workflow establishes an authoritative daily settlement baseline, converting granular wire events into consolidated counterparty liquidity measures required for financial auditability.",
                    "inputs": [
                        {
                            "source_dataset": "wire_transfers.csv",
                            "business_role": "Inbound interbank wire transfer log feed",
                            "source_format": "CSV Data File",
                            "dependency_significance": "Essential transactional input required for central ledger reconciliation"
                        }
                    ],
                    "outputs": [
                        {
                            "output_deliverable": "daily_settlement.xlsx",
                            "what_it_represents": "Finalized institutional settlement position matrix",
                            "business_use": "Daily regulatory liquidity reporting and central bank clearing verification",
                            "destination_format": "Excel Workbook"
                        }
                    ],
                    "sequential_stages": [
                        {
                            "stage_number": 1,
                            "stage_name": "Transaction Ingestion & Aggregation",
                            "description": "Ingests raw wire transfer events and computes gross settlement exposure",
                            "operational_explanation": "Processes wire transfer records via aggregation"
                        }
                    ],
                    "business_rules": [
                        {
                            "business_rule": "Aggregate Settlement Exposure",
                            "category": "Aggregation",
                            "evidence_configuration": "SUM(Transfer_Amount), COUNT DISTINCT(Account_ID)"
                        }
                    ],
                    "lineage": [
                        {
                            "source_datasets": "wire_transfers.csv",
                            "major_business_transformation": "Aggregate transfer totals and count distinct accounts",
                            "target_deliverable": "daily_settlement.xlsx"
                        }
                    ]
                })
            elif "sales_invoices.xlsx" in user or "Sales Commission" in user:
                return json.dumps({
                    "workflow_title": "Commercial Sales Commission Reconciliation",
                    "workflow_description": "Computes sales team commission accruals from completed customer invoices.",
                    "executive_summary": "This business report evaluates commercial sales revenue and associate commission calculation logic. The workflow ingests finalized customer sales invoices, calculates contractual 5% commission amounts using derived formulas, and generates the monthly commission reconciliation ledger.",
                    "methods_of_analysis": "The workflow applies mathematical formula expressions ([Invoice_Total] * 0.05) to derive associate payout metrics from transaction-level invoice data.",
                    "findings": [
                        "Contractual commission accruals are calculated using deterministic formula derivation ([Invoice_Total] * 0.05), ensuring standard commercial rate application across all completed invoices.",
                        "Inbound invoice data is validated against sales ledger constraints prior to commission ledger posting.",
                        "The resulting commission ledger is published to sales_commission_ledger.csv for monthly payroll audit and revenue operations review."
                    ],
                    "conclusions": "The workflow establishes a standardized commercial commission derivation pipeline, transforming raw invoice data into audited compensation records.",
                    "inputs": [
                        {
                            "source_dataset": "sales_invoices.xlsx",
                            "business_role": "Completed customer sales invoice register",
                            "source_format": "Excel Workbook",
                            "dependency_significance": "Primary revenue source feed required for commission calculation"
                        }
                    ],
                    "outputs": [
                        {
                            "output_deliverable": "sales_commission_ledger.csv",
                            "what_it_represents": "Itemized sales representative commission register",
                            "business_use": "Monthly commercial sales payroll processing and revenue audit",
                            "destination_format": "CSV Data File"
                        }
                    ],
                    "sequential_stages": [
                        {
                            "stage_number": 1,
                            "stage_name": "Invoice Ingestion & Commission Derivation",
                            "description": "Ingests sales invoices and computes commission accruals",
                            "operational_explanation": "Applies formula derivation to invoices"
                        }
                    ],
                    "business_rules": [
                        {
                            "business_rule": "Calculate 5% Commission",
                            "category": "Calculation",
                            "evidence_configuration": "[Invoice_Total] * 0.05"
                        }
                    ],
                    "lineage": [
                        {
                            "source_datasets": "sales_invoices.xlsx",
                            "major_business_transformation": "Calculate commission accruals from invoice totals",
                            "target_deliverable": "sales_commission_ledger.csv"
                        }
                    ]
                })
            elif "policy_claims_register.xlsx" in user or "Claims Loss Frequency" in user:
                return json.dumps({
                    "workflow_title": "Insurance Claims Loss Frequency Analysis",
                    "workflow_description": "Aggregates quarterly policyholder loss amounts and claim frequency by underwriting class.",
                    "executive_summary": "This actuarial business report evaluates policyholder incurred claim losses and loss frequency. The workflow ingests the centralized policy claims register, computes aggregate incurred loss amounts via multi-dimensional summation, enumerates distinct policy exposure counts, and publishes the finalized loss ratio matrix for quarterly underwriting review.",
                    "methods_of_analysis": "The methodology implements actuarial loss summation (SUM on Incurred_Loss) and distinct policy enumeration (COUNT DISTINCT on Policy_Number) to establish policyholder loss experience profiles.",
                    "findings": [
                        "Incurred claims exposure is aggregated using SUM(Incurred_Loss), converting transactional loss events into quarterly actuarial portfolio totals.",
                        "Policy exposure frequency is measured via COUNT DISTINCT(Policy_Number), determining active underwriting participation.",
                        "The resulting actuarial loss matrix is exported to quarterly_actuarial_losses.xlsx for recurring reserve review and risk pricing adjustments."
                    ],
                    "conclusions": "The workflow establishes an actuarial reporting baseline, converting transactional loss events into structured portfolio frequency and severity metrics.",
                    "inputs": [
                        {
                            "source_dataset": "policy_claims_register.xlsx",
                            "business_role": "Central policyholder claims register",
                            "source_format": "Excel Workbook",
                            "dependency_significance": "Authoritative actuarial loss register for quarterly risk assessment"
                        }
                    ],
                    "outputs": [
                        {
                            "output_deliverable": "quarterly_actuarial_losses.xlsx",
                            "what_it_represents": "Quarterly underwriting loss ratio and claim severity matrix",
                            "business_use": "Quarterly reserve audit and actuarial portfolio risk review",
                            "destination_format": "Excel Workbook"
                        }
                    ],
                    "sequential_stages": [
                        {
                            "stage_number": 1,
                            "stage_name": "Claims Ingestion & Actuarial Aggregation",
                            "description": "Ingests policy claims and computes portfolio loss aggregations",
                            "operational_explanation": "Aggregates policyholder claims records"
                        }
                    ],
                    "business_rules": [
                        {
                            "business_rule": "Aggregate Portfolio Incurred Losses",
                            "category": "Aggregation",
                            "evidence_configuration": "SUM(Incurred_Loss), COUNT DISTINCT(Policy_Number)"
                        }
                    ],
                    "lineage": [
                        {
                            "source_datasets": "policy_claims_register.xlsx",
                            "major_business_transformation": "Aggregate incurred losses and count unique policyholders",
                            "target_deliverable": "quarterly_actuarial_losses.xlsx"
                        }
                    ]
                })
            return None

        client = FakeLLMClient(generator_fn=mock_domain_llm)
        gen = LLMNarrativeGenerator(client=client, cache=LLMNarrativeCache())
        set_default_llm_client(client)
        set_default_generator(gen)

        try:
            # 1. Generate Banking Report
            b_graph = build_graph(banking_wf)
            b_order = execution_order(b_graph)
            b_bs = generate_business_summary(banking_wf, b_graph, b_order)
            b_rep = gen.generate_business_report(banking_wf, b_bs, graph=b_graph, workflow_id="banking-1")
            assert b_rep is not None

            b_bs.business_purpose = b_rep.executive_summary
            b_bs.executive_summary.subject_and_purpose = b_rep.executive_summary
            b_bs.executive_summary.methods_and_process = b_rep.methods_of_analysis
            b_bs.executive_summary.findings = b_rep.findings
            b_bs.executive_summary.conclusions = b_rep.conclusions

            b_doc_model = build_document_model(
                workflow=banking_wf,
                execution_order=b_order,
                translations={},
                dag_layout=None,
                lineage_paths=[],
                business_summary=b_bs,
                analysis_id="banking-1",
            )
            b_docx = tmp_path / "banking_report.docx"
            generate_docx(b_doc_model, b_docx)
            assert b_docx.exists()
            b_text = self._extract_all_docx_text(b_docx)

            # 2. Generate Sales Report
            s_graph = build_graph(sales_wf)
            s_order = execution_order(s_graph)
            s_bs = generate_business_summary(sales_wf, s_graph, s_order)
            s_rep = gen.generate_business_report(sales_wf, s_bs, graph=s_graph, workflow_id="sales-1")
            assert s_rep is not None

            s_bs.business_purpose = s_rep.executive_summary
            s_bs.executive_summary.subject_and_purpose = s_rep.executive_summary
            s_bs.executive_summary.methods_and_process = s_rep.methods_of_analysis
            s_bs.executive_summary.findings = s_rep.findings
            s_bs.executive_summary.conclusions = s_rep.conclusions

            s_doc_model = build_document_model(
                workflow=sales_wf,
                execution_order=s_order,
                translations={},
                dag_layout=None,
                lineage_paths=[],
                business_summary=s_bs,
                analysis_id="sales-1",
            )
            s_docx = tmp_path / "sales_report.docx"
            generate_docx(s_doc_model, s_docx)
            assert s_docx.exists()
            s_text = self._extract_all_docx_text(s_docx)

            # 3. Generate Claims Report
            c_graph = build_graph(claims_wf)
            c_order = execution_order(c_graph)
            c_bs = generate_business_summary(claims_wf, c_graph, c_order)
            c_rep = gen.generate_business_report(claims_wf, c_bs, graph=c_graph, workflow_id="claims-1")
            assert c_rep is not None

            c_bs.business_purpose = c_rep.executive_summary
            c_bs.executive_summary.subject_and_purpose = c_rep.executive_summary
            c_bs.executive_summary.methods_and_process = c_rep.methods_of_analysis
            c_bs.executive_summary.findings = c_rep.findings
            c_bs.executive_summary.conclusions = c_rep.conclusions

            c_doc_model = build_document_model(
                workflow=claims_wf,
                execution_order=c_order,
                translations={},
                dag_layout=None,
                lineage_paths=[],
                business_summary=c_bs,
                analysis_id="claims-1",
            )
            c_docx = tmp_path / "claims_loss_report.docx"
            generate_docx(c_doc_model, c_docx)
            assert c_docx.exists()
            c_text = self._extract_all_docx_text(c_docx)

            # 4. Verify Banking-specific terms and absence of Sales / Claims terms
            assert "wire_transfers.csv" in b_text
            assert "daily_settlement.xlsx" in b_text
            assert "liquidity" in b_text.lower()
            assert "sales_invoices.xlsx" not in b_text
            assert "commission" not in b_text.lower()
            assert "policy_claims_register.xlsx" not in b_text
            assert "actuarial" not in b_text.lower()

            # 5. Verify Sales-specific terms and absence of Banking / Claims terms
            assert "sales_invoices.xlsx" in s_text
            assert "sales_commission_ledger.csv" in s_text
            assert "commission" in s_text.lower()
            assert "wire_transfers.csv" not in s_text
            assert "settlement" not in s_text.lower()
            assert "policy_claims_register.xlsx" not in s_text
            assert "actuarial" not in s_text.lower()

            # 6. Verify Claims-specific terms and absence of Banking / Sales terms
            assert "policy_claims_register.xlsx" in c_text
            assert "quarterly_actuarial_losses.xlsx" in c_text
            assert "actuarial" in c_text.lower()
            assert "wire_transfers.csv" not in c_text
            assert "settlement" not in c_text.lower()
            assert "sales_invoices.xlsx" not in c_text
            assert "commission" not in c_text.lower()

            # 7. Verify absolute absence of prohibited sections in all 3 documents
            for doc_text in (b_text, s_text, c_text):
                assert "Recommendations" not in doc_text
                assert "Limitations" not in doc_text
                assert "Visual Workflow Graph" not in doc_text
                assert "Visual DAG" not in doc_text
                assert "DAG Architecture" not in doc_text
                assert "Section 5" not in doc_text
                assert "5. Visual" not in doc_text

        finally:
            set_default_llm_client(None)
            set_default_generator(None)




"""Tests for DOCX and SVG output generators."""

from pathlib import Path
import docx
import pytest

from awa.parser.xml_parser import parse_workflow
from awa.graph.builder import build_graph, execution_order, build_input_map
from awa.graph.dag_layouter import compute_dag_layout
from awa.graph.lineage import compute_lineage_paths
from awa.translators.registry import get_translator
import awa.translators  # register translators
from awa.generators.doc_builder import build_document_model
from awa.generators.docx_generator import generate_docx
from awa.generators.svg_generator import generate_svg


def test_docx_generation(tmp_path: Path):
    wf = parse_workflow("fixtures/joins/join_workflow.yxmd")
    g = build_graph(wf)
    order = execution_order(g)
    input_map = build_input_map(wf)

    translations = {}
    for tid in order:
        tool = wf.tools[tid]
        tr = get_translator(tool)
        translations[tid] = tr.translate(tool, input_map.get(tid, []), wf)

    layout = compute_dag_layout(g, wf, order)
    lineage = compute_lineage_paths(wf, g)

    doc_model = build_document_model(wf, order, translations, layout, lineage)
    assert "join_workflow" in doc_model.title
    assert len(doc_model.nodes) == 6
    assert len(doc_model.execution_order) == 6

    docx_path = tmp_path / "workflow.docx"
    generate_docx(doc_model, docx_path)

    assert docx_path.exists()
    assert docx_path.stat().st_size > 0

    # Verify python-docx can open and read it
    doc = docx.Document(str(docx_path))
    headings = [p.text for p in doc.paragraphs if p.text.startswith("1.") or p.text.startswith("2.") or p.text.startswith("3.")]
    assert len(headings) >= 3
    assert len(doc.tables) >= 3

    # Assert no forbidden support level words exist in document text
    full_text = "\n".join(p.text for p in doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            full_text += "\n" + " | ".join(c.text for c in row.cells)

    forbidden = [
        "Support Level", "FULL", "SUPPORTED", "PARTIAL", "PASS-THROUGH",
        "PASS_THROUGH", "DOCUMENTATION_ONLY", "EXTERNAL_EXECUTION",
        "UNSUPPORTED", "Analysis Diagnostics"
    ]
    for word in forbidden:
        assert word not in full_text, f"Forbidden word '{word}' found in DOCX!"

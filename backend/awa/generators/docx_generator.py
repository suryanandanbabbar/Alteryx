"""Business-oriented DOCX report generator for AWA.

Generates a publication-quality Executive Business Report (.docx)
with a structured Executive Summary answering:
1. PURPOSE
2. INPUT -> PROCESS -> OUTPUT
3. KEY BUSINESS LOGIC
4. BUSINESS USE / INTERPRETATION
5. KEY CONSIDERATIONS
All sections are strictly conditional and never render empty headings.

Followed by detailed visual DAG graph, step-by-step tool specifications,
and technical configuration appendix.

LLM-free and 100% deterministic.
"""

from __future__ import annotations

import io
from pathlib import Path
from typing import Any
from PIL import Image, ImageDraw

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls

from awa.model.doc_model import DocumentModel
from awa.model.dag_layout import DagLayout
from awa.model.visual_category import get_category_colors


# Brand Color Palette (Corporate Consulting Style)
COLOR_NAVY_HEX = "0f172a"
COLOR_PRIMARY_HEX = "ea580c"
COLOR_SECONDARY_HEX = "1e293b"
COLOR_BORDER_HEX = "cbd5e1"
COLOR_BG_LIGHT_HEX = "f8fafc"
COLOR_MUTED_HEX = "64748b"

RGB_NAVY = RGBColor(15, 23, 42)
RGB_PRIMARY = RGBColor(234, 88, 12)
RGB_MUTED = RGBColor(100, 116, 139)
RGB_DARK = RGBColor(30, 41, 59)
RGB_WHITE = RGBColor(255, 255, 255)


def set_cell_background(cell: Any, fill_hex: str) -> None:
    """Set background color of a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="{fill_hex}"/>')
    tcPr.append(shd)


def set_cell_margins(cell: Any, top: int = 80, bottom: int = 80, left: int = 120, right: int = 120) -> None:
    """Set inner cell padding in dxa (1 pt = 20 dxa)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in (('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)):
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)


def render_dag_to_png(layout: DagLayout, scale: float = 1.0) -> bytes:
    """Render the DagLayout into a high-resolution PNG image for DOCX embedding."""
    width = int(max(layout.width, 400.0) * scale)
    height = int(max(layout.height + 60.0, 260.0) * scale)

    img = Image.new("RGB", (width, height), color="#0d1326")
    draw = ImageDraw.Draw(img)

    # Header Text
    draw.text(
        (int(20 * scale), int(16 * scale)),
        f"Alteryx Workflow DAG — {layout.title}",
        fill="#f8fafc",
    )
    draw.text(
        (int(20 * scale), int(34 * scale)),
        f"{len(layout.nodes)} workflow tools · {len(layout.edges)} data connections",
        fill="#94a3b8",
    )

    y_offset = int(45 * scale)

    # Draw Edges
    for edge in layout.edges:
        if not edge.path_points or len(edge.path_points) < 2:
            continue
        pts = [(int(p[0] * scale), int(p[1] * scale + y_offset)) for p in edge.path_points]
        has_label = bool(edge.source_anchor and edge.source_anchor not in ("Output", "0", "#1"))
        edge_color = "#38bdf8" if has_label else "#64748b"

        for i in range(len(pts) - 1):
            draw.line([pts[i], pts[i + 1]], fill=edge_color, width=max(1, int(1.5 * scale)))

    # Draw Nodes
    for node in layout.nodes:
        cat_color = get_category_colors(node.visual_category)
        nx = int(node.x * scale)
        ny = int(node.y * scale + y_offset)
        nw = int(node.width * scale)
        nh = int(node.height * scale)

        # Card body
        draw.rounded_rectangle(
            [nx, ny, nx + nw, ny + nh],
            radius=int(6 * scale),
            fill=cat_color["fill"],
            outline=cat_color["stroke"],
            width=max(1, int(1.5 * scale)),
        )

        # Badge circle
        cx = nx + int(20 * scale)
        cy = ny + int(nh / 2)
        r = int(10 * scale)
        draw.ellipse(
            [cx - r, cy - r, cx + r, cy + r],
            fill=cat_color["stroke"],
            outline=cat_color["stroke"],
        )

        id_str = str(node.tool_id)
        draw.text(
            (cx - int(4 * scale), cy - int(6 * scale)),
            id_str,
            fill="#0a0f1d",
        )

        # Tool Type Text
        type_display = node.tool_type if len(node.tool_type) <= 16 else node.tool_type[:14] + ".."
        draw.text(
            (nx + int(36 * scale), cy - int(10 * scale)),
            type_display,
            fill=cat_color["text"],
        )

        # Tool Name Subtitle
        label_display = node.label if len(node.label) <= 18 else node.label[:16] + ".."
        draw.text(
            (nx + int(36 * scale), cy + int(4 * scale)),
            label_display,
            fill="#94a3b8",
        )

    buf = io.BytesIO()
    img.save(buf, format="PNG")
    return buf.getvalue()


def generate_docx(
    doc_model: DocumentModel,
    output_path: Path | str,
    svg_content: str | None = None,
) -> None:
    """Generate workflow.docx with a structured, professional Executive Summary.

    Args:
        doc_model: Canonical format-independent documentation model.
        output_path: Target path for the .docx file.
        svg_content: Optional SVG string.
    """
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()

    # Configure Margins (1 inch all around)
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    workflow_name = doc_model.metadata.get("Workflow Name", "Alteryx Workflow")
    bs = doc_model.business_summary
    exec_summary = bs.executive_summary if bs else None

    # -------------------------------------------------------------
    # Cover Section / Document Header Banner
    # -------------------------------------------------------------
    p_brand = doc.add_paragraph()
    p_brand.paragraph_format.space_before = Pt(0)
    p_brand.paragraph_format.space_after = Pt(2)
    run_brand = p_brand.add_run("AWA — ALTERYX WORKFLOW ANALYZER")
    run_brand.font.size = Pt(9)
    run_brand.font.bold = True
    run_brand.font.color.rgb = RGB_PRIMARY

    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_before = Pt(0)
    p_title.paragraph_format.space_after = Pt(4)
    run_title = p_title.add_run(f"EXECUTIVE BUSINESS REPORT\n{workflow_name}")
    run_title.font.size = Pt(20)
    run_title.font.bold = True
    run_title.font.color.rgb = RGB_NAVY

    p_subtitle = doc.add_paragraph()
    p_subtitle.paragraph_format.space_after = Pt(14)
    one_line = bs.one_line_purpose if bs and bs.one_line_purpose else "Data preparation and business reporting workflow"
    run_sub = p_subtitle.add_run(one_line)
    run_sub.font.size = Pt(11)
    run_sub.font.color.rgb = RGB_MUTED

    # =============================================================
    # SECTION 1: EXECUTIVE SUMMARY (Business Analysis Assessment)
    # =============================================================
    h1 = doc.add_heading("1. Executive Summary", level=1)
    h1.paragraph_format.space_before = Pt(8)
    h1.paragraph_format.space_after = Pt(6)

    # 1.1 Subject Matter / Business Purpose (Conditional)
    purpose_text = exec_summary.subject_and_purpose if exec_summary and exec_summary.subject_and_purpose else (bs.business_purpose if bs else "")
    if purpose_text:
        p_purp = doc.add_paragraph()
        p_purp.paragraph_format.space_after = Pt(8)
        r_p = p_purp.add_run(purpose_text)
        r_p.font.size = Pt(9.5)

    # 1.2 Methods of Analysis / Workflow Process (Conditional)
    if exec_summary and exec_summary.methods_and_process:
        p_meth_hdr = doc.add_paragraph()
        p_meth_hdr.paragraph_format.space_before = Pt(4)
        p_meth_hdr.paragraph_format.space_after = Pt(2)
        r_mh = p_meth_hdr.add_run("Methods of Analysis")
        r_mh.bold = True
        r_mh.font.color.rgb = RGB_NAVY
        r_mh.font.size = Pt(10)

        p_meth = doc.add_paragraph()
        p_meth.paragraph_format.space_after = Pt(8)
        r_m = p_meth.add_run(exec_summary.methods_and_process)
        r_m.font.size = Pt(9.5)

    # 1.3 Findings (Conditional)
    if exec_summary and exec_summary.findings:
        p_find_hdr = doc.add_paragraph()
        p_find_hdr.paragraph_format.space_before = Pt(4)
        p_find_hdr.paragraph_format.space_after = Pt(2)
        r_fh = p_find_hdr.add_run("Findings")
        r_fh.bold = True
        r_fh.font.color.rgb = RGB_NAVY
        r_fh.font.size = Pt(10)

        for finding in exec_summary.findings:
            p_f = doc.add_paragraph(style='List Bullet')
            p_f.paragraph_format.space_after = Pt(2)
            r_f = p_f.add_run(finding)
            r_f.font.size = Pt(9.0)

        doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # 1.4 Conclusions (Conditional)
    if exec_summary and exec_summary.conclusions:
        p_conc_hdr = doc.add_paragraph()
        p_conc_hdr.paragraph_format.space_before = Pt(4)
        p_conc_hdr.paragraph_format.space_after = Pt(2)
        r_ch = p_conc_hdr.add_run("Conclusions")
        r_ch.bold = True
        r_ch.font.color.rgb = RGB_NAVY
        r_ch.font.size = Pt(10)

        p_conc = doc.add_paragraph()
        p_conc.paragraph_format.space_after = Pt(8)
        r_c = p_conc.add_run(exec_summary.conclusions)
        r_c.font.size = Pt(9.5)

    # 1.5 Recommendations (Conditional)
    if exec_summary and exec_summary.recommendations:
        p_rec_hdr = doc.add_paragraph()
        p_rec_hdr.paragraph_format.space_before = Pt(4)
        p_rec_hdr.paragraph_format.space_after = Pt(2)
        r_rh = p_rec_hdr.add_run("Recommendations")
        r_rh.bold = True
        r_rh.font.color.rgb = RGB_NAVY
        r_rh.font.size = Pt(10)

        for rec in exec_summary.recommendations:
            p_r = doc.add_paragraph(style='List Bullet')
            p_r.paragraph_format.space_after = Pt(2)
            r_r = p_r.add_run(rec)
            r_r.font.size = Pt(9.0)

    # =============================================================
    # SECTION 2: BUSINESS PROCESS & OPERATIONAL DELIVERABLES
    # =============================================================
    has_process_section = bool(bs and (bs.source_inputs or bs.business_outputs or bs.processing_stages))
    if has_process_section:
        h2 = doc.add_heading("2. Business Process & Operational Deliverables", level=1)
        h2.paragraph_format.space_before = Pt(16)
        h2.paragraph_format.space_after = Pt(6)

        # 2.1 Inputs & Dependencies
        if bs and bs.source_inputs:
            p_in_tbl_hdr = doc.add_paragraph()
            p_in_tbl_hdr.paragraph_format.space_after = Pt(3)
            r_ith = p_in_tbl_hdr.add_run("2.1 Inputs & Upstream Dependencies")
            r_ith.bold = True
            r_ith.font.color.rgb = RGB_NAVY
            r_ith.font.size = Pt(10)

            in_table = doc.add_table(rows=1, cols=4)
            in_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            in_hdr = in_table.rows[0].cells
            in_hdr[0].text = "Source Dataset"
            in_hdr[1].text = "Business Role"
            in_hdr[2].text = "Source Format"
            in_hdr[3].text = "Dependency Significance"

            for c in in_hdr:
                set_cell_background(c, COLOR_SECONDARY_HEX)
                c.paragraphs[0].runs[0].font.bold = True
                c.paragraphs[0].runs[0].font.color.rgb = RGB_WHITE
                c.paragraphs[0].runs[0].font.size = Pt(8.0)
                set_cell_margins(c, top=60, bottom=60, left=80, right=80)

            for inp in bs.source_inputs:
                row = in_table.add_row().cells
                set_cell_background(row[0], COLOR_BG_LIGHT_HEX)
                for c in row:
                    set_cell_margins(c, top=50, bottom=50, left=80, right=80)

                p0 = row[0].paragraphs[0]
                p0.add_run(inp.name).bold = True
                p0.runs[0].font.size = Pt(8.0)

                p1 = row[1].paragraphs[0]
                p1.add_run(inp.business_role or inp.description or "Source input stream")
                p1.runs[0].font.size = Pt(8.0)

                p2 = row[2].paragraphs[0]
                p2.add_run(inp.source_type)
                p2.runs[0].font.size = Pt(8.0)

                p3 = row[3].paragraphs[0]
                p3.add_run(getattr(inp, "dependency_significance", "Primary source input required for downstream processing"))
                p3.runs[0].font.size = Pt(8.0)

            doc.add_paragraph().paragraph_format.space_after = Pt(8)

        # 2.2 Outputs & Business Use
        if bs and bs.business_outputs:
            p_out_tbl_hdr = doc.add_paragraph()
            p_out_tbl_hdr.paragraph_format.space_after = Pt(3)
            r_oth = p_out_tbl_hdr.add_run("2.2 Outputs & Business Reporting Deliverables")
            r_oth.bold = True
            r_oth.font.color.rgb = RGB_NAVY
            r_oth.font.size = Pt(10)

            out_table = doc.add_table(rows=1, cols=4)
            out_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            out_hdr = out_table.rows[0].cells
            out_hdr[0].text = "Output Deliverable"
            out_hdr[1].text = "What it Represents"
            out_hdr[2].text = "Business Use"
            out_hdr[3].text = "Destination Format"

            for c in out_hdr:
                set_cell_background(c, COLOR_NAVY_HEX)
                c.paragraphs[0].runs[0].font.bold = True
                c.paragraphs[0].runs[0].font.color.rgb = RGB_WHITE
                c.paragraphs[0].runs[0].font.size = Pt(8.0)
                set_cell_margins(c, top=60, bottom=60, left=80, right=80)

            for out in bs.business_outputs:
                row = out_table.add_row().cells
                set_cell_background(row[0], COLOR_BG_LIGHT_HEX)
                for c in row:
                    set_cell_margins(c, top=50, bottom=50, left=80, right=80)

                p0 = row[0].paragraphs[0]
                p0.add_run(out.name).bold = True
                p0.runs[0].font.size = Pt(8.0)

                p1 = row[1].paragraphs[0]
                p1.add_run(out.business_meaning or out.business_purpose)
                p1.runs[0].font.size = Pt(8.0)

                p2 = row[2].paragraphs[0]
                use_val = out.likely_use if out.likely_use and out.likely_use != "Use not documented" else "Business use: Not documented"
                p2.add_run(use_val)
                p2.runs[0].font.size = Pt(8.0)
                if "Not documented" in use_val:
                    p2.runs[0].font.italic = True
                    p2.runs[0].font.color.rgb = RGB_MUTED

                p3 = row[3].paragraphs[0]
                p3.add_run(out.destination_type)
                p3.runs[0].font.size = Pt(8.0)

            doc.add_paragraph().paragraph_format.space_after = Pt(8)

        # 2.3 Operational Stages
        if bs and bs.processing_stages:
            p_stg_hdr = doc.add_paragraph()
            p_stg_hdr.paragraph_format.space_after = Pt(3)
            r_stgh = p_stg_hdr.add_run("2.3 Sequential Operational Stages")
            r_stgh.bold = True
            r_stgh.font.color.rgb = RGB_NAVY
            r_stgh.font.size = Pt(10)

            for stg in bs.processing_stages:
                p_stg = doc.add_paragraph()
                p_stg.paragraph_format.left_indent = Inches(0.2)
                p_stg.paragraph_format.space_after = Pt(3)
                r_num = p_stg.add_run(f"Stage {stg.stage_number:02d} — {stg.name}: ")
                r_num.bold = True
                r_num.font.size = Pt(8.5)
                r_num.font.color.rgb = RGB_NAVY

                r_desc = p_stg.add_run(f"{stg.summary or stg.description}. ")
                r_desc.font.size = Pt(8.5)

                if stg.major_transformation:
                    r_t = p_stg.add_run(f"({stg.major_transformation})")
                    r_t.font.size = Pt(8.0)
                    r_t.font.color.rgb = RGB_MUTED
                    r_t.font.italic = True

            doc.add_paragraph().paragraph_format.space_after = Pt(16)

    # =============================================================
    # SECTION 3: KEY BUSINESS RULES & TRANSFORMATIONS
    # =============================================================
    if bs and bs.business_rules:
        h3 = doc.add_heading("3. Key Business Rules & Transformations", level=1)
        h3.paragraph_format.space_before = Pt(16)
        h3.paragraph_format.space_after = Pt(6)

        doc.add_paragraph(
            "The following business rules and operational derivations have been extracted directly from workflow tool configurations and annotations:"
        )

        rules_table = doc.add_table(rows=1, cols=3)
        rules_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        r_hdr = rules_table.rows[0].cells
        r_hdr[0].text = "Business Rule"
        r_hdr[1].text = "Category"
        r_hdr[2].text = "Evidence / Configuration"

        for c in r_hdr:
            set_cell_background(c, COLOR_SECONDARY_HEX)
            c.paragraphs[0].runs[0].font.bold = True
            c.paragraphs[0].runs[0].font.color.rgb = RGB_WHITE
            c.paragraphs[0].runs[0].font.size = Pt(8.0)
            set_cell_margins(c, top=60, bottom=60, left=80, right=80)

        for rule in bs.business_rules:
            row = rules_table.add_row().cells
            set_cell_background(row[0], COLOR_BG_LIGHT_HEX)
            for c in row:
                set_cell_margins(c, top=50, bottom=50, left=80, right=80)

            p0 = row[0].paragraphs[0]
            p0.add_run(rule.description).bold = True
            p0.runs[0].font.size = Pt(8.0)

            p1 = row[1].paragraphs[0]
            p1.add_run(rule.category)
            p1.runs[0].font.size = Pt(8.0)

            p2 = row[2].paragraphs[0]
            p2.add_run(rule.evidence)
            p2.runs[0].font.size = Pt(8.0)

        doc.add_paragraph().paragraph_format.space_after = Pt(16)

    # =============================================================
    # SECTION 4: SOURCE-TO-TARGET DATA LINEAGE
    # =============================================================
    if bs and bs.lineage:
        h4 = doc.add_heading("4. Source-to-Target Data Lineage", level=1)
        h4.paragraph_format.space_before = Pt(16)
        h4.paragraph_format.space_after = Pt(6)

        doc.add_paragraph(
            "The following matrix maps source datasets through intermediate transformation operations to finalized published deliverables:"
        )

        lin_table = doc.add_table(rows=1, cols=3)
        lin_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        l_hdr = lin_table.rows[0].cells
        l_hdr[0].text = "Source Dataset(s)"
        l_hdr[1].text = "Major Business Transformation"
        l_hdr[2].text = "Target Deliverable"

        for c in l_hdr:
            set_cell_background(c, COLOR_NAVY_HEX)
            c.paragraphs[0].runs[0].font.bold = True
            c.paragraphs[0].runs[0].font.color.rgb = RGB_WHITE
            c.paragraphs[0].runs[0].font.size = Pt(8.0)
            set_cell_margins(c, top=60, bottom=60, left=80, right=80)

        for lin in bs.lineage:
            row = lin_table.add_row().cells
            set_cell_background(row[0], COLOR_BG_LIGHT_HEX)
            for c in row:
                set_cell_margins(c, top=50, bottom=50, left=80, right=80)

            p0 = row[0].paragraphs[0]
            p0.add_run(lin.source_name).bold = True
            p0.runs[0].font.size = Pt(8.0)

            p1 = row[1].paragraphs[0]
            p1.add_run(lin.transformation or lin.transformation_summary)
            p1.runs[0].font.size = Pt(8.0)

            p2 = row[2].paragraphs[0]
            p2.add_run(lin.target_name).bold = True
            p2.runs[0].font.size = Pt(8.0)

        doc.add_paragraph().paragraph_format.space_after = Pt(16)

    # =============================================================
    # SECTION 5: VISUAL WORKFLOW GRAPH (DAG ARCHITECTURE)
    # =============================================================
    h5 = doc.add_heading("5. Visual Workflow Graph (DAG Architecture)", level=1)
    h5.paragraph_format.space_before = Pt(16)
    h5.paragraph_format.space_after = Pt(6)

    doc.add_paragraph(
        "The following diagram illustrates the complete directed acyclic graph (DAG) representing all "
        "data flows, transformation nodes, and execution branches across the workflow:"
    )

    try:
        png_bytes = render_dag_to_png(doc_model.dag_layout, scale=1.5)
        image_stream = io.BytesIO(png_bytes)
        doc.add_picture(image_stream, width=Inches(6.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception as e:
        doc.add_paragraph(f"[Graph rendering note: DAG preview could not be generated: {e}]")

    doc.add_paragraph().paragraph_format.space_after = Pt(16)

    # =============================================================
    # SECTION 6: STEP-BY-STEP TOOL SPECIFICATIONS
    # =============================================================
    h6 = doc.add_heading("6. Step-by-Step Tool Specifications", level=1)
    h6.paragraph_format.space_before = Pt(16)
    h6.paragraph_format.space_after = Pt(6)

    doc.add_paragraph(
        "This section documents each workflow execution step in topological sequence, detailing the "
        "business action and operational function of every individual tool:"
    )

    for step in doc_model.execution_order:
        matching_node = next((n for n in doc_model.nodes if n.tool_id == step.tool_id), None)
        annotation_text = matching_node.annotation if matching_node and matching_node.annotation else ""
        technical_summary = step.summary or (matching_node.summary if matching_node else "Processes records in the workflow pipeline.")

        # Business action priority: Annotation or synthesized action
        business_action = annotation_text or f"Executes {step.tool_type} operation in the data pipeline."

        p_step = doc.add_paragraph()
        p_step.paragraph_format.space_before = Pt(8)
        p_step.paragraph_format.space_after = Pt(2)
        r_num = p_step.add_run(f"Step {step.step_number:02d}: ")
        r_num.bold = True
        r_num.font.color.rgb = RGB_PRIMARY
        r_num.font.size = Pt(9.5)

        r_name = p_step.add_run(f"{step.name} ")
        r_name.bold = True
        r_name.font.color.rgb = RGB_NAVY
        r_name.font.size = Pt(9.5)

        r_type = p_step.add_run(f"({step.tool_type})")
        r_type.font.color.rgb = RGB_MUTED
        r_type.font.size = Pt(8.5)

        if step.container_name:
            r_cont = p_step.add_run(f"  [Container: {step.container_name}]")
            r_cont.font.color.rgb = RGB_MUTED
            r_cont.font.size = Pt(8.0)
            r_cont.font.italic = True

        p_desc = doc.add_paragraph()
        p_desc.paragraph_format.left_indent = Inches(0.25)
        p_desc.paragraph_format.space_after = Pt(2)
        r_ba = p_desc.add_run("Business Action: ")
        r_ba.bold = True
        r_ba.font.size = Pt(8.5)
        r_bat = p_desc.add_run(business_action)
        r_bat.font.size = Pt(8.5)

        p_tech = doc.add_paragraph()
        p_tech.paragraph_format.left_indent = Inches(0.25)
        p_tech.paragraph_format.space_after = Pt(4)
        r_ta = p_tech.add_run("Technical Function: ")
        r_ta.italic = True
        r_ta.font.size = Pt(8.0)
        r_ta.font.color.rgb = RGB_MUTED
        r_tat = p_tech.add_run(technical_summary)
        r_tat.font.size = Pt(8.0)
        r_tat.font.color.rgb = RGB_MUTED

    doc.add_paragraph().paragraph_format.space_after = Pt(16)

    # =============================================================
    # SECTION 7: TECHNICAL CONFIGURATION APPENDIX
    # =============================================================
    h7 = doc.add_heading("7. Technical Configuration Appendix", level=1)
    h7.paragraph_format.space_before = Pt(16)
    h7.paragraph_format.space_after = Pt(6)

    doc.add_paragraph(
        "This appendix documents the parsed configuration parameters and connection strings for each workflow tool for engineering review:"
    )

    for node in doc_model.nodes:
        p_node_title = doc.add_paragraph()
        p_node_title.paragraph_format.space_before = Pt(10)
        p_node_title.paragraph_format.space_after = Pt(4)
        cont_str = f" [Container: {node.container_name}]" if node.container_name else ""
        r_nh = p_node_title.add_run(f"Tool #{node.tool_id} — {node.name} ({node.tool_type}){cont_str}")
        r_nh.bold = True
        r_nh.font.size = Pt(9.0)
        r_nh.font.color.rgb = RGB_NAVY

        # Configuration Table
        if node.configuration:
            cfg_table = doc.add_table(rows=1, cols=2)
            cfg_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            c_hdr = cfg_table.rows[0].cells
            c_hdr[0].text = "Parameter"
            c_hdr[1].text = "Value"
            set_cell_background(c_hdr[0], COLOR_SECONDARY_HEX)
            set_cell_background(c_hdr[1], COLOR_SECONDARY_HEX)
            for c in c_hdr:
                c.paragraphs[0].runs[0].font.bold = True
                c.paragraphs[0].runs[0].font.color.rgb = RGB_WHITE
                c.paragraphs[0].runs[0].font.size = Pt(8.0)
                set_cell_margins(c, top=50, bottom=50, left=80, right=80)

            for key, val in node.configuration.items():
                r_cells = cfg_table.add_row().cells
                set_cell_background(r_cells[0], COLOR_BG_LIGHT_HEX)
                for c in r_cells:
                    set_cell_margins(c, top=50, bottom=50, left=80, right=80)

                pk = r_cells[0].paragraphs[0]
                pk.add_run(str(key))
                pk.runs[0].font.size = Pt(8.0)
                pk.runs[0].bold = True

                pv = r_cells[1].paragraphs[0]
                val_str = str(val) if not isinstance(val, (dict, list)) else str(val)[:120]
                pv.add_run(val_str)
                pv.runs[0].font.size = Pt(8.0)

    doc.save(str(output_path))

"""DOCX documentation generator — produces workflow.docx.

Renders a rich, professional Word document from the canonical DocumentModel,
including an embedded high-resolution visual DAG diagram.
"""

from __future__ import annotations

import io
import math
from pathlib import Path
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from awa.model.doc_model import DocumentModel
from awa.model.dag_layout import DagLayout
from awa.model.visual_category import get_category_colors


def render_dag_image(layout: DagLayout, scale: float = 2.0) -> bytes:
    """Render the canonical DagLayout as a high-resolution raster PNG image.

    Directly consumes the canonical DagLayout geometry so the diagram in DOCX
    matches workflow.svg and the React DAG viewer exactly.

    Args:
        layout: Canonical DAG layout.
        scale: Resolution multiplier for sharp print quality (default 2.0).

    Returns:
        PNG image bytes.
    """
    width = int(max(layout.width, 500.0) * scale)
    height = int(max(layout.height + 70.0, 280.0) * scale)

    img = Image.new("RGB", (width, height), color="#0d1326")
    draw = ImageDraw.Draw(img)

    # 1. Header Title
    title_text = f"Alteryx Workflow DAG — {layout.title}"
    sub_text = f"{len(layout.nodes)} nodes · {len(layout.edges)} connections"

    draw.text(
        (int(24 * scale), int(20 * scale)),
        title_text,
        fill="#f8fafc",
    )
    draw.text(
        (int(24 * scale), int(40 * scale)),
        sub_text,
        fill="#64748b",
    )

    offset_y = 50.0 * scale

    # 2. Render Edges with Arrowheads
    for edge in layout.edges:
        if not edge.path_points or len(edge.path_points) < 2:
            continue

        pts = [(int(p[0] * scale), int(p[1] * scale + offset_y)) for p in edge.path_points]
        has_label = bool(edge.source_anchor and edge.source_anchor not in ("Output", "0", "#1"))
        edge_color = "#38bdf8" if has_label else "#475569"

        # Draw line segments along waypoints
        for i in range(len(pts) - 1):
            draw.line([pts[i], pts[i + 1]], fill=edge_color, width=max(2, int(2 * scale)))

        # Draw arrowhead at the end point
        end_pt = pts[-1]
        prev_pt = pts[-2]
        dx = end_pt[0] - prev_pt[0]
        dy = end_pt[1] - prev_pt[1]
        angle = math.atan2(dy, dx)
        arrow_len = 10.0 * scale
        arrow_width = 5.0 * scale

        arrow_p1 = (
            int(end_pt[0] - arrow_len * math.cos(angle) + arrow_width * math.sin(angle)),
            int(end_pt[1] - arrow_len * math.sin(angle) - arrow_width * math.cos(angle)),
        )
        arrow_p2 = (
            int(end_pt[0] - arrow_len * math.cos(angle) - arrow_width * math.sin(angle)),
            int(end_pt[1] - arrow_len * math.sin(angle) + arrow_width * math.cos(angle)),
        )
        draw.polygon([end_pt, arrow_p1, arrow_p2], fill=edge_color)

        # Draw edge label badge if branching
        if has_label and len(pts) >= 2:
            mid_x = (pts[0][0] + pts[-1][0]) // 2
            mid_y = (pts[0][1] + pts[-1][1]) // 2 - int(8 * scale)
            badge_w = int(36 * scale)
            badge_h = int(14 * scale)
            draw.rounded_rectangle(
                [
                    mid_x - badge_w // 2,
                    mid_y - badge_h // 2,
                    mid_x + badge_w // 2,
                    mid_y + badge_h // 2,
                ],
                radius=int(3 * scale),
                fill="#0f172a",
                outline="#38bdf8",
                width=1,
            )
            draw.text(
                (mid_x - int(10 * scale), mid_y - int(5 * scale)),
                edge.source_anchor,
                fill="#38bdf8",
            )

    # 3. Render Nodes
    for node in layout.nodes:
        cat_color = get_category_colors(node.visual_category)

        nx = int(node.x * scale)
        ny = int(node.y * scale + offset_y)
        nw = int(node.width * scale)
        nh = int(node.height * scale)

        # Outer rounded box
        draw.rounded_rectangle(
            [nx, ny, nx + nw, ny + nh],
            radius=int(8 * scale),
            fill=cat_color["fill"],
            outline=cat_color["stroke"],
            width=max(2, int(1.5 * scale)),
        )

        # Tool ID circle badge
        cx = nx + int(24 * scale)
        cy = ny + nh // 2
        r = int(12 * scale)
        draw.ellipse(
            [cx - r, cy - r, cx + r, cy + r],
            fill=cat_color["stroke"],
            outline=None,
        )

        # Tool ID text
        id_str = str(node.tool_id)
        draw.text(
            (cx - int(4 * scale), cy - int(6 * scale)),
            id_str,
            fill="#0a0f1d",
        )

        # Tool Type Text
        type_display = node.tool_type if len(node.tool_type) <= 16 else node.tool_type[:14] + ".."
        draw.text(
            (nx + int(44 * scale), cy - int(10 * scale)),
            type_display,
            fill=cat_color["text"],
        )

        # Tool Name Subtitle
        label_display = node.label if len(node.label) <= 18 else node.label[:16] + ".."
        draw.text(
            (nx + int(44 * scale), cy + int(4 * scale)),
            label_display,
            fill="#94a3b8",
        )

    buf = io.BytesIO()
    img.save(buf, format="PNG")
    return buf.getvalue()


def generate_docx(
    doc_model: DocumentModel,
    output_path: Path | str,
    svg_content: str | None = None,
) -> None:
    """Generate workflow.docx from a DocumentModel with an embedded DAG diagram.

    Args:
        doc_model: Canonical format-independent documentation model.
        output_path: Target path for the .docx file.
        svg_content: Optional SVG string (for fallback).
    """
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()

    # --- Document Header ---
    title = doc.add_heading(doc_model.title, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    p = doc.add_paragraph()
    p.add_run("Automated Deterministic Workflow Analysis & Lineage Report").italic = True

    # 1. Executive Summary & Metadata
    doc.add_heading("1. Workflow Metadata & Summary", level=1)
    
    meta_table = doc.add_table(rows=1, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = meta_table.rows[0].cells
    hdr_cells[0].text = "Property"
    hdr_cells[1].text = "Value"

    for prop, val in doc_model.metadata.items():
        row_cells = meta_table.add_row().cells
        row_cells[0].text = str(prop)
        row_cells[1].text = str(val)

    # Metrics Table
    p_metrics = doc.add_paragraph()
    p_metrics.add_run("\nWorkflow Metrics:").bold = True

    metrics_table = doc.add_table(rows=1, cols=2)
    metrics_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    m_hdr = metrics_table.rows[0].cells
    m_hdr[0].text = "Metric"
    m_hdr[1].text = "Count"

    for metric_name, count in doc_model.metrics.items():
        row_cells = metrics_table.add_row().cells
        row_cells[0].text = str(metric_name)
        row_cells[1].text = str(count)

    # 2. Embedded Visual DAG Diagram
    doc.add_heading("2. Workflow Graph (Visual DAG)", level=1)
    if doc_model.dag_layout.nodes:
        dag_png_bytes = render_dag_image(doc_model.dag_layout)
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_picture(io.BytesIO(dag_png_bytes), width=Inches(6.0))
        p_caption = doc.add_paragraph()
        p_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_caption.add_run(f"Figure 1: Visual DAG for {doc_model.dag_layout.title}").italic = True
    else:
        doc.add_paragraph("No nodes present in workflow graph.")

    # 3. Execution Order
    doc.add_heading("3. Execution Order (Topological Sort)", level=1)
    doc.add_paragraph(
        "Tools execute strictly in data-flow dependency order as determined by graph topology:"
    )

    exec_table = doc.add_table(rows=1, cols=4)
    exec_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    e_hdr = exec_table.rows[0].cells
    e_hdr[0].text = "Step #"
    e_hdr[1].text = "Tool ID"
    e_hdr[2].text = "Type"
    e_hdr[3].text = "Name / Annotation"

    for step in doc_model.execution_order:
        row_cells = exec_table.add_row().cells
        row_cells[0].text = str(step.step_number)
        row_cells[1].text = str(step.tool_id)
        row_cells[2].text = str(step.tool_type)
        row_cells[3].text = str(step.name)

    # 4. Node-by-Node Details
    doc.add_heading("4. Detailed Node Configurations", level=1)

    for node in doc_model.nodes:
        doc.add_heading(f"Tool #{node.tool_id} — {node.name} ({node.tool_type})", level=2)
        
        info_p = doc.add_paragraph()
        info_p.add_run(f"Plugin: {node.plugin}\n")
        info_p.add_run(f"Support Level: ").bold = True
        info_p.add_run(f"{node.support_level.upper()}\n")
        if node.description:
            info_p.add_run(f"Description: {node.description}\n")

        if node.input_variables:
            info_p.add_run(f"Input Variables: {', '.join(node.input_variables)}\n")
        if node.output_variables:
            out_str = ", ".join(f"{k} → {v}" for k, v in node.output_variables.items())
            info_p.add_run(f"Output Variables: {out_str}\n")

        # Configuration Table
        if node.configuration:
            config_table = doc.add_table(rows=1, cols=2)
            c_hdr = config_table.rows[0].cells
            c_hdr[0].text = "Setting"
            c_hdr[1].text = "Value"
            for k, v in sorted(node.configuration.items()):
                val_str = str(v) if not isinstance(v, (dict, list)) else str(v)[:100]
                row_cells = config_table.add_row().cells
                row_cells[0].text = str(k)
                row_cells[1].text = val_str

        # Tool Diagnostics
        if node.diagnostics:
            diag_p = doc.add_paragraph()
            diag_p.add_run("Diagnostics:").bold = True
            for d in node.diagnostics:
                diag_p.add_run(f"\n• [{d.level.value.upper()}] {d.message}")

    # 5. Data Lineage
    doc.add_heading("5. Data Lineage Paths", level=1)
    if doc_model.lineage_paths:
        for idx, lp in enumerate(doc_model.lineage_paths, start=1):
            path_str = " → ".join(
                f"{name} (#{tid})" for tid, name in zip(lp.tool_ids, lp.tool_names)
            )
            doc.add_paragraph(f"{idx}. {path_str}")
    else:
        doc.add_paragraph("No source-to-sink lineage paths detected.")

    # 6. External Dependencies
    doc.add_heading("6. External Dependencies", level=1)
    if doc_model.dependencies:
        dep_table = doc.add_table(rows=1, cols=3)
        d_hdr = dep_table.rows[0].cells
        d_hdr[0].text = "Type"
        d_hdr[1].text = "Reference"
        d_hdr[2].text = "Associated Tool"
        for dep in doc_model.dependencies:
            row_cells = dep_table.add_row().cells
            row_cells[0].text = str(dep.dep_type)
            row_cells[1].text = str(dep.reference)
            row_cells[2].text = f"Tool #{dep.tool_id}" if dep.tool_id else "Workflow"
    else:
        doc.add_paragraph("No external file, database, or macro dependencies detected.")

    # 7. Global Diagnostics & Unsupported Features
    doc.add_heading("7. Analysis Diagnostics", level=1)
    if doc_model.diagnostics:
        for diag in doc_model.diagnostics:
            tool_info = f" (Tool #{diag.tool_id})" if diag.tool_id else ""
            p_diag = doc.add_paragraph(style="List Bullet")
            p_diag.add_run(f"[{diag.level.value.upper()}]{tool_info}: ").bold = True
            p_diag.add_run(diag.message)
            if diag.detail:
                p_diag.add_run(f" — {diag.detail}").italic = True
    else:
        doc.add_paragraph("No warnings or errors detected during analysis.")

    doc.save(str(output_path))
